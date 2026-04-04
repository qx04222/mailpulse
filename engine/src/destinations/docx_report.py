"""
生成类 Manus 质量的结构化 DOCX 周报。
输入：Sonnet 返回的结构化 JSON dict
输出：DOCX bytes
"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from datetime import datetime

# 状态颜色映射
STATUS_COLORS = {
    "new_inquiry": RGBColor(220, 53, 69),    # 红 — 新询价
    "quoting": RGBColor(255, 193, 7),         # 黄 — 报价中
    "quoted": RGBColor(40, 167, 69),          # 绿 — 已报价
    "negotiating": RGBColor(255, 153, 0),     # 橙 — 谈判中
    "follow_up": RGBColor(23, 162, 184),      # 青 — 待跟进
    "closed": RGBColor(108, 117, 125),        # 灰 — 已关闭
    "info_only": RGBColor(108, 117, 125),     # 灰 — 仅信息
}

PRIORITY_LABELS = {
    "high": ("🔴 高", RGBColor(220, 53, 69)),
    "medium": ("🟡 中", RGBColor(255, 153, 0)),
    "low": ("🟢 低", RGBColor(40, 167, 69)),
}

# 中文字体
CJK_FONT = "Microsoft YaHei"  # DOCX 内嵌字体名，兼容 Windows/macOS/Linux


def _set_cell_shading(cell, color_hex: str):
    """设置单元格背景色"""
    shading = cell._element.get_or_add_tcPr()
    shading_elem = shading.makeelement(qn('w:shd'), {
        qn('w:fill'): color_hex,
        qn('w:val'): 'clear',
    })
    shading.append(shading_elem)


def _add_styled_paragraph(doc, text: str, font_size: int = 10, bold: bool = False,
                           color: RGBColor = None, alignment=None, space_after: int = 6):
    """添加带样式的段落"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(font_size)
    run.font.name = CJK_FONT
    run._element.rPr.rFonts.set(qn('w:eastAsia'), CJK_FONT)
    if bold:
        run.bold = True
    if color:
        run.font.color.rgb = color
    if alignment:
        p.alignment = alignment
    p.paragraph_format.space_after = Pt(space_after)
    return p


def _create_table(doc, headers: list, rows: list, col_widths: list = None,
                  header_color: str = "2E75B6") -> None:
    """创建带样式的表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # 表头
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(header)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.name = CJK_FONT
        run._element.rPr.rFonts.set(qn('w:eastAsia'), CJK_FONT)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_cell_shading(cell, header_color)

    # 数据行
    for row_idx, row_data in enumerate(rows):
        for col_idx, cell_text in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(cell_text))
            run.font.size = Pt(9)
            run.font.name = CJK_FONT
            run._element.rPr.rFonts.set(qn('w:eastAsia'), CJK_FONT)
            # 交替行背景
            if row_idx % 2 == 1:
                _set_cell_shading(cell, "F2F2F2")

    # 列宽
    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(width)

    doc.add_paragraph()  # 表后间距


def _report_label(date_range: str) -> str:
    """Determine '日报' vs '周报' from date_range string."""
    try:
        parts = date_range.replace("–", "-").split("-")
        if len(parts) == 2:
            start_md = parts[0].strip().split("/")
            end_md = parts[1].strip().split("/")
            start_day = int(start_md[1]) if len(start_md) == 2 else int(start_md[0])
            end_day = int(end_md[1]) if len(end_md) == 2 else int(end_md[0])
            if abs(end_day - start_day) <= 1:
                return "邮件日报"
    except Exception:
        pass
    return "邮件日报" if "–" not in date_range and "-" not in date_range else "邮件周报"


def generate_report_docx(digest_data: dict, company_name: str, date_range: str) -> bytes:
    """生成结构化 DOCX 报告"""
    doc = Document()

    # 设置默认字体
    style = doc.styles['Normal']
    style.font.name = CJK_FONT
    style.font.size = Pt(10)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), CJK_FONT)

    overview = digest_data.get("overview", {})
    label = _report_label(date_range)

    # ── 标题页 ──────────────────────────────────────────────────
    _add_styled_paragraph(doc, f"{company_name} {label}",
                          font_size=22, bold=True,
                          color=RGBColor(46, 117, 182),
                          alignment=WD_ALIGN_PARAGRAPH.CENTER,
                          space_after=4)
    _add_styled_paragraph(doc, f"统计周期：{date_range}",
                          font_size=11,
                          color=RGBColor(108, 117, 125),
                          alignment=WD_ALIGN_PARAGRAPH.CENTER,
                          space_after=2)
    _add_styled_paragraph(doc, f"生成时间：{datetime.now().strftime('%Y年%m月%d日 %H:%M')}",
                          font_size=9,
                          color=RGBColor(150, 150, 150),
                          alignment=WD_ALIGN_PARAGRAPH.CENTER,
                          space_after=12)

    # ── 一、总览 ────────────────────────────────────────────────
    doc.add_heading('一、总览', level=1)

    highlights = overview.get("highlights", "")
    if highlights:
        _add_styled_paragraph(doc, highlights, font_size=10, space_after=8)

    # 人员统计表
    per_person = overview.get("per_person_stats", [])
    if per_person:
        headers = ["报价人", "客户数量", "已报价", "待处理/跟进", "已解决"]
        rows = []
        for p in per_person:
            rows.append([
                p.get("name", ""),
                str(p.get("client_count", 0)),
                str(p.get("quoted", 0)),
                str(p.get("pending", 0)),
                str(p.get("resolved", 0)),
            ])
        # 合计行
        totals = ["合计",
                  str(sum(p.get("client_count", 0) for p in per_person)),
                  str(sum(p.get("quoted", 0) for p in per_person)),
                  str(sum(p.get("pending", 0) for p in per_person)),
                  str(sum(p.get("resolved", 0) for p in per_person))]
        rows.append(totals)
        _create_table(doc, headers, rows, col_widths=[4, 2.5, 2.5, 3, 2.5])

    # ── 二、客户详情 ──────────────────────────────────────────
    clients = digest_data.get("clients", [])
    if clients:
        doc.add_heading('二、按客户的详细情况', level=1)

        # 按负责人分组
        by_person = {}
        for c in clients:
            person = c.get("assigned_to", "待分配")
            if person not in by_person:
                by_person[person] = []
            by_person[person].append(c)

        client_num = 1
        for person, person_clients in by_person.items():
            doc.add_heading(f'{person} 负责的客户（{len(person_clients)}个）', level=2)

            for c in person_clients:
                # 客户标题
                status_label = c.get("status_label", "")
                status = c.get("status", "info_only")
                status_color = STATUS_COLORS.get(status, RGBColor(108, 117, 125))

                doc.add_heading(f'{client_num}. {c.get("client_name", "未知客户")}', level=3)

                # 客户信息表
                info_rows = []
                if c.get("contact_email"):
                    info_rows.append(["联系方式", c["contact_email"]])
                if c.get("project_address"):
                    info_rows.append(["项目地址", c["project_address"]])
                if c.get("product_type"):
                    info_rows.append(["产品类型", c["product_type"]])
                info_rows.append(["当前状态", f"{status_label}"])
                info_rows.append(["邮件数量", str(c.get("email_count", 0))])
                if c.get("latest_date"):
                    info_rows.append(["最新日期", c["latest_date"]])

                if info_rows:
                    _create_table(doc, ["项目", "详情"], info_rows, col_widths=[3.5, 11])

                # 详细描述
                summary = c.get("summary", "")
                if summary:
                    _add_styled_paragraph(doc, summary, font_size=10, space_after=4)

                # 关键要点
                key_details = c.get("key_details", [])
                if key_details:
                    for detail in key_details:
                        p = doc.add_paragraph(style='List Bullet')
                        run = p.add_run(detail)
                        run.font.size = Pt(9)
                        run.font.name = CJK_FONT
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), CJK_FONT)

                # 行动建议
                action = c.get("action_needed")
                if action:
                    _add_styled_paragraph(doc, f"⚠ 行动建议：{action}",
                                          font_size=10, bold=True,
                                          color=RGBColor(220, 53, 69),
                                          space_after=8)

                client_num += 1

    # ── 三、跟进状态汇总 ──────────────────────────────────────
    followup = digest_data.get("followup_update", {})
    has_followup = any([
        followup.get("resolved"),
        followup.get("overdue"),
        followup.get("still_pending"),
    ])

    if has_followup:
        doc.add_heading('三、跟进状态汇总', level=1)

        # 汇总表
        status_rows = []
        resolved = followup.get("resolved", [])
        overdue = followup.get("overdue", [])
        pending = followup.get("still_pending", [])

        resolved_names = ", ".join(r.get("subject", "")[:20] for r in resolved) if resolved else ""
        overdue_names = ", ".join(o.get("subject", "")[:20] for o in overdue) if overdue else ""
        pending_names = ", ".join(p.get("subject", "")[:20] for p in pending) if pending else ""

        summary_headers = ["状态", "数量", "详情"]
        summary_rows = []
        if resolved:
            summary_rows.append(["✅ 已解决", str(len(resolved)), resolved_names])
        if overdue:
            summary_rows.append(["🚨 超期未处理", str(len(overdue)), overdue_names])
        if pending:
            summary_rows.append(["⏳ 持续跟进", str(len(pending)), pending_names])

        if summary_rows:
            _create_table(doc, summary_headers, summary_rows, col_widths=[3.5, 2, 9])

    # ── 四、需立即处理的事项 ──────────────────────────────────
    actions = digest_data.get("priority_actions", [])
    if actions:
        doc.add_heading('四、需要立即处理的事项（按优先级排列）', level=1)

        for priority_level in ["high", "medium", "low"]:
            level_actions = [a for a in actions if a.get("priority") == priority_level]
            if not level_actions:
                continue

            label, color = PRIORITY_LABELS.get(priority_level, ("", RGBColor(0, 0, 0)))
            _add_styled_paragraph(doc, f"{label}优先级", font_size=12, bold=True, color=color)

            for a in level_actions:
                client_info = f"（{a['client']}）" if a.get("client") else ""
                assigned_info = f" — {a['assigned_to']}" if a.get("assigned_to") else ""
                _add_styled_paragraph(
                    doc,
                    f"• {a['action']}{client_info}{assigned_info}",
                    font_size=10,
                    space_after=4,
                )
                if a.get("deadline"):
                    _add_styled_paragraph(
                        doc, f"  建议截止：{a['deadline']}",
                        font_size=9, color=RGBColor(108, 117, 125),
                        space_after=2,
                    )

    # ── 五、垃圾箱/删除审查 ──────────────────────────────────
    trash = digest_data.get("trash_spam_review", [])
    worth_checking = [t for t in trash if t.get("worth_checking")]
    if worth_checking:
        doc.add_heading('五、垃圾箱/删除邮件审查', level=1)
        _add_styled_paragraph(doc, "以下被删除/标记垃圾的邮件可能需要关注：", font_size=10)
        for t in worth_checking:
            _add_styled_paragraph(
                doc,
                f"• [{t.get('bucket', 'trash').upper()}] {t['sender']} — {t['subject']}\n  {t.get('reason', '')}",
                font_size=9, space_after=4,
            )

    # ── 页脚信息 ──────────────────────────────────────────────
    doc.add_paragraph()
    _add_styled_paragraph(
        doc,
        f"文档生成时间：{datetime.now().strftime('%Y年%m月%d日')}\n数据来源：Gmail 自动分析 · Claude AI",
        font_size=8,
        color=RGBColor(150, 150, 150),
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
    )

    # 输出 bytes
    from io import BytesIO
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
